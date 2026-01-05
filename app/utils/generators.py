from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re

def add_hyperlink(paragraph, text, url, is_bold=False):
    """
    Add a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add color (Blue)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rPr.append(color)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    if is_bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    new_run.append(rPr)
    
    # Add text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)

    # Append hyperlink to the paragraph element
    paragraph._p.append(hyperlink)

    return hyperlink

def process_text(paragraph, text, default_bold=False, default_color=None, font_size=None):
    """
    Process text for bold markers (**text**) and markdown links ([text](url)).
    """
    # Regex to find **bold** OR [link](url)
    # This regex captures delimiters to keep them in the split list
    pattern = r'(\*\*.*?\*\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            content = part[2:-2]
            # Check if there is a link inside the bold text
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', content)
            if link_match:
                # Link inside bold
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                add_hyperlink(paragraph, link_text, link_url, is_bold=True)
            else:
                run = paragraph.add_run(content)
                run.bold = True
                if default_color:
                    run.font.color.rgb = default_color
                if font_size:
                    run.font.size = font_size
                    
        elif part.startswith('[') and part.endswith(')') and '](' in part:
            # Link
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                add_hyperlink(paragraph, link_text, link_url, is_bold=default_bold)
            else:
                # Fallback
                run = paragraph.add_run(part)
                if default_bold: run.bold = True
                if default_color: run.font.color.rgb = default_color
                if font_size: run.font.size = font_size
        else:
            # Normal text
            run = paragraph.add_run(part)
            if default_bold: run.bold = True
            if default_color: run.font.color.rgb = default_color
            if font_size: run.font.size = font_size

def create_docx(content: str) -> BytesIO:
    doc = Document()
    
    # Set default font to Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    
    # Set narrow margins for compact look
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # State tracking for centering header info
    first_h1_seen = False
    waiting_for_contact_info = False

    # Simple Markdown Parser
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            text = line[2:].replace('**', '')
            p = doc.add_heading(text, level=1)
            
            # Check if this is the first H1 (Candidate Name)
            if not first_h1_seen:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Increase font size for the name
                for run in p.runs:
                    run.font.size = Pt(24)
                    run.font.name = 'Calibri'
                
                first_h1_seen = True
                waiting_for_contact_info = True
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                waiting_for_contact_info = False

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            
        elif line.startswith('## '):
            # Heading 2
            waiting_for_contact_info = False
            text = line[3:].replace('**', '')
            p = doc.add_heading(text, level=2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith('### '):
            # Heading 3
            waiting_for_contact_info = False
            text = line[4:].replace('**', '')
            p = doc.add_heading(text, level=3)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        elif '|' in line and not line.startswith('- ') and not line.startswith('* '):
            waiting_for_contact_info = False
            # Handle Date Alignment (Role | Date)
            # Split on the LAST pipe to handle cases like "Role | Location | Date"
            last_pipe_index = line.rfind('|')
            
            if last_pipe_index != -1:
                left_part = line[:last_pipe_index].strip()
                right_part = line[last_pipe_index+1:].strip()
                
                # Check if it's a valid Role/Education line (shouldn't be too long, e.g. > 120 chars)
                if len(left_part) < 120:
                    table = doc.add_table(rows=1, cols=2)
                    table.autofit = False
                    table.allow_autofit = False
                    
                    # Calculate widths in twips (1440 twips per inch)
                    # Page is 8.5", margins are 0.5" each, so content width is 7.5"
                    total_width = int(7.5 * 1440)  # 10800 twips
                    left_width = int(5.625 * 1440)  # 75% = 8100 twips
                    right_width = int(1.875 * 1440)  # 25% = 2700 twips
                    
                    # Access the table XML element
                    tbl = table._tbl
                    
                    # Get or create tblPr (table properties)
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    # Remove any existing width/layout settings to avoid conflicts
                    for child in list(tblPr):
                        if child.tag.endswith('tblW') or child.tag.endswith('tblLayout'):
                            tblPr.remove(child)
                    
                    # Set table width
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), str(total_width))
                    tblW.set(qn('w:type'), 'dxa')
                    tblPr.append(tblW)
                    
                    # Set fixed table layout
                    tblLayout = OxmlElement('w:tblLayout')
                    tblLayout.set(qn('w:type'), 'fixed')
                    tblPr.append(tblLayout)
                    
                    # Remove existing tblGrid if present and create new one
                    existing_grid = tbl.find(qn('w:tblGrid'))
                    if existing_grid is not None:
                        tbl.remove(existing_grid)
                    
                    # Create table grid with explicit column widths
                    tblGrid = OxmlElement('w:tblGrid')
                    gridCol1 = OxmlElement('w:gridCol')
                    gridCol1.set(qn('w:w'), str(left_width))
                    gridCol2 = OxmlElement('w:gridCol')
                    gridCol2.set(qn('w:w'), str(right_width))
                    tblGrid.append(gridCol1)
                    tblGrid.append(gridCol2)
                    # Insert tblGrid after tblPr
                    tbl.insert(1, tblGrid)
                    
                    # Set cell widths explicitly
                    left_cell = table.cell(0, 0)
                    left_cell.width = Inches(5.625)
                    tc1 = left_cell._tc
                    tcPr1 = tc1.get_or_add_tcPr()
                    # Remove any existing tcW
                    for child in list(tcPr1):
                        if child.tag.endswith('tcW'):
                            tcPr1.remove(child)
                    tcW1 = OxmlElement('w:tcW')
                    tcW1.set(qn('w:w'), str(left_width))
                    tcW1.set(qn('w:type'), 'dxa')
                    tcPr1.insert(0, tcW1)
                    
                    right_cell = table.cell(0, 1)
                    right_cell.width = Inches(1.875)
                    tc2 = right_cell._tc
                    tcPr2 = tc2.get_or_add_tcPr()
                    # Remove any existing tcW
                    for child in list(tcPr2):
                        if child.tag.endswith('tcW'):
                            tcPr2.remove(child)
                    tcW2 = OxmlElement('w:tcW')
                    tcW2.set(qn('w:w'), str(right_width))
                    tcW2.set(qn('w:type'), 'dxa')
                    tcPr2.insert(0, tcW2)
                    
                    # Prevent table from being split across pages
                    cantSplit = OxmlElement('w:cantSplit')
                    for tr in tbl.tr_lst:
                        trPr = tr.get_or_add_trPr()
                        trPr.append(cantSplit)

                    # Left Cell (Role)
                    left_cell = table.cell(0, 0)
                    left_p = left_cell.paragraphs[0]
                    
                    # Handle bold in left cell (Standard Black)
                    # Clear default paragraph content
                    left_p.clear()
                    process_text(left_p, left_part)

                    # Right Cell (Date)
                    right_cell = table.cell(0, 1)
                    right_p = right_cell.paragraphs[0]
                    right_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                    # Handle bold in right cell
                    right_p.clear()
                    process_text(right_p, right_part)

                else:
                    # Fallback if too long
                    p = doc.add_paragraph()
                    process_text(p, line)
            else:
                # Fallback if split fails
                p = doc.add_paragraph()
                process_text(p, line)
        elif line.startswith('- ') or line.startswith('* '):
            waiting_for_contact_info = False
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            process_text(p, line[2:], font_size=Pt(10))
        else:
            # Normal text
            p = doc.add_paragraph()
            
            # Check if this is the contact info line
            if waiting_for_contact_info:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                waiting_for_contact_info = False  # Only center the first paragraph after name
            
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            process_text(p, line, font_size=Pt(10))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
